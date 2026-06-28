"""
Genera el Short Paper TP1 — Sinapsistencia (formato Scitepress/WEBIST).
Uso:
  python generate_short_paper.py          # inglés (por defecto)
  python generate_short_paper.py --lang es
  python generate_short_paper.py --lang en --lang es  # ambos
"""
import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUT_DIR = Path(__file__).parent

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(10)
HEADING_SIZE = Pt(11)

CONTENT = {
    "en": {
        "suffix": "",
        "title": (
            "Sinapsistencia: A Web Platform for Medical-Legal Mediation "
            "with Explainable Machine Learning in the Peruvian Health Sector"
        ),
        "affiliation": (
            "\u00b9School of Systems and Computer Engineering, Peruvian University of Applied Sciences, "
            "2390 Primavera Avenue, Lima, Peru"
        ),
        "keywords_label": "Keywords:",
        "keywords": (
            "Medical-Legal Mediation, Machine Learning, Explainable AI, "
            "Random Forest, TF-IDF, Health Informatics."
        ),
        "abstract_label": "Abstract:",
        "abstract": (
            "Healthcare professionals in Peru face growing exposure to medical malpractice claims and "
            "often lack structured, timely access to specialized legal advice. Informal channels "
            "delay case management and reduce traceability. This study presents Sinapsistencia, a "
            "web platform that connects physicians with medical-law attorneys through a microservices "
            "architecture (Angular 21, Spring Boot 3.5, PostgreSQL, and a FastAPI machine learning "
            "service). The system supports medical-legal case registration with simulated clinical "
            "context compliant with Peruvian data protection law (Law 29733), document traceability, "
            "role-based dashboards, and two explainable ML components: a Random Forest model for legal "
            "risk assessment and a TF-IDF cosine-similarity matcher for attorney recommendation. "
            "The risk model was trained on 6,000 synthetic cases using a rule-based label generator "
            "and achieved macro precision of 0.618, recall of 0.597, and F1-score of 0.603 on a "
            "held-out validation set. Functional validation included five automated integration test "
            "suites and incremental execution of acceptance scenarios. Average inference latency for "
            "risk assessment remained below 200 ms in local experiments. Results indicate that a "
            "hybrid ML approach with explicit XAI factors and cold-start fallbacks is feasible for "
            "academic MVP validation in a sensitive domain where AI must act as decision support rather "
            "than autonomous judgment."
        ),
        "sections": {
            "intro": "1 INTRODUCTION",
            "related": "2 RELATED WORKS",
            "solution": "3 PROPOSED SOLUTION",
            "experiments": "4 EXPERIMENTS AND VALIDATION",
            "results": "5 RESULTS",
            "conclusions": "6 CONCLUSIONS",
            "ack": "ACKNOWLEDGMENTS",
            "refs": "REFERENCES",
        },
        "intro": [
            "Medical malpractice and negligence disputes in Peru are increasingly visible through "
            "institutional supervision and patient advocacy channels. SUSALUD reported more than "
            "163,000 health-rights attentions in 2024 alone, reflecting sustained demand for formal "
            "recourse mechanisms. Physicians, however, frequently rely on informal networks to obtain "
            "legal orientation, with estimated response delays of 48 to 72 hours and limited audit "
            "trails (Miziara & Miziara, 2022; Tumelty, 2023).",
            "This problem is relevant because delayed legal guidance increases professional vulnerability, "
            "encourages defensive medicine, and hampers early case structuring (Gunenc et al., 2024; "
            "Rowland et al., 2022). Digital platforms could improve mediation, yet naive approaches fail "
            "for three reasons: (i) medical-legal data are sensitive and regulated; (ii) black-box "
            "recommendations are ethically unacceptable in liability contexts (Jones et al., 2023; "
            "Scott et al., 2024); and (iii) cold-start conditions limit collaborative filtering when "
            "historical attorney-case data are scarce.",
            "Prior solutions address fragments of the problem—electronic records, generic legal marketplaces, "
            "or clinical decision support—but not integrated physician-to-attorney mediation with "
            "explainable prioritization under role-based governance. Sinapsistencia differs by combining "
            "(1) simulated non-identifiable case context, (2) microservice separation of ML inference, "
            "(3) persisted XAI explanations, and (4) specialty-based fallback matching when the ML "
            "service is offline.",
            "The key components of our approach are: a three-portal web application (physician, attorney, "
            "administrator); a Spring Boot API with ownership enforcement replacing database-level RLS; "
            "a FastAPI service implementing Random Forest risk scoring and TF-IDF attorney ranking; and "
            "an audit trail for critical actions. Current limitations include synthetic training data, "
            "pilot validation without real patient records, and pending large-scale functional test "
            "execution against a 90-case academic test plan.",
            "The remainder of this paper is organized as follows: Section 2 reviews related work; "
            "Section 3 describes the proposed architecture and algorithms; Section 4 explains the "
            "validation methodology; Section 5 presents results; and Section 6 concludes with "
            "future work.",
        ],
        "rw_intro": (
            "Research on medical-legal risk, AI-assisted decision support, and recommender systems "
            "provides the foundation for Sinapsistencia. We group prior efforts into three lines."
        ),
        "related": [
            (
                "Medical negligence and digital liability.",
                "Miziara and Miziara (2022) review medical errors, negligence, and defensive medicine, "
                "highlighting documentation and communication as recurring failure points. Rowland et al. "
                "(2022) analyze digital health-specific malpractice risks, emphasizing traceability and "
                "accountability. Wong et al. (2021) examine emergency department claims, showing that "
                "procedural and documentation issues dominate litigation patterns. These studies justify "
                "structured case intake and document versioning but do not provide integrated mediation "
                "workflows between clinicians and attorneys.",
            ),
            (
                "Trustworthy AI in clinical and legal contexts.",
                "Jones et al. (2023) report clinicians' concerns about trust and liability when using "
                "AI decision support. Wang et al. (2023) systematically review human-centered design for "
                "AI-enabled clinical systems, recommending transparency and explicit limitations. Scott "
                "et al. (2024) further note adoption barriers when explanations are absent. Our platform "
                "addresses this gap by persisting risk factors derived from feature importance and by "
                "displaying non-deterministic advisory messages in the user interface.",
            ),
            (
                "Matching and risk modeling techniques.",
                "Manning et al. (2008) and Ricci et al. (2015) establish TF-IDF and content-based "
                "recommendation as effective cold-start strategies—relevant for attorney matching without "
                "dense interaction logs. Random Forest classifiers remain a robust baseline for tabular "
                "risk scoring with interpretable feature contributions (Aggarwal, 2016). Unlike prior "
                "legal-tech marketplaces, Sinapsistencia couples both techniques in a health-specific "
                "workflow with simulated data aligned to Peruvian privacy constraints (Congreso de la "
                "República del Perú, 2011).",
            ),
        ],
        "solution_intro": (
            "Sinapsistencia follows a decoupled microservices architecture. The Angular 21 frontend "
            "consumes REST APIs from a Spring Boot 3.5 backend (Java 21) secured with JWT and role-based "
            "access control. PostgreSQL 16 stores operational data and Flyway migrations version the "
            "schema. A dedicated FastAPI service hosts ML models, invoked by the backend through an "
            "HTTP proxy with a 5-second timeout and graceful degradation."
        ),
        "solution_subs": [
            (
                "3.1 System Architecture",
                "The physician portal supports case creation with simulated legal-clinical context "
                "(reference age, medical area, summary, relevant factors) without identifiable patient data. "
                "The attorney portal lists relevant pending cases and contact requests. The administrator "
                "portal manages users, audit logs, ML metrics, and summary reports. Ownership rules ensure "
                "that physicians cannot access peers' cases—enforced in the service layer and validated by "
                "automated integration tests.",
            ),
            (
                "3.2 Risk Assessment Model (Random Forest)",
                "The risk module predicts risk_score (0–1) and risk_level (bajo, moderado, alto, critico) "
                "from structured features: medical specialty, procedure complexity, priority, documentation "
                "completeness, informed consent, prior complaints, and days since incident. Training data "
                "comprise 6,000 synthetic rows generated by business rules plus Gaussian noise to avoid "
                "trivial memorization. A RandomForestClassifier (200 trees) and RandomForestRegressor share "
                "the same preprocessed feature space (one-hot specialty, ordinal complexity/priority). "
                "Explainability combines global feature importance with normalized case-specific values to "
                "produce risk_factors shown to users as advisory context.",
            ),
            (
                "3.3 Attorney Matching (TF-IDF + Cosine Similarity)",
                "Attorney profiles are vectorized with TfidfVectorizer over legal specialties, medical areas "
                "of interest, and biographical text. Physician queries built from specialty and case context "
                "are compared via cosine similarity. Results include content_score, matched specialties, and "
                "natural-language reasons. When the ML service is unavailable, the backend applies a cold-start "
                "fallback ranking attorneys by medical-area overlap.",
            ),
            (
                "3.4 API Contract",
                "FastAPI exposes GET /health, GET /api/v1/model/info, POST /api/v1/risk-assessment, and "
                "POST /api/v1/recommendations. Spring Boot proxies these endpoints and normalizes responses "
                "to camelCase for the frontend. High/critical risk levels trigger optional fire-and-forget "
                "notifications to an external automation webhook.",
            ),
        ],
        "experiments_intro": (
            "Validation followed four tracks aligned with the project quality plan: (1) ML offline "
            "evaluation, (2) backend integration testing, (3) frontend build verification, and "
            "(4) incremental functional scenarios from the academic test catalog."
        ),
        "experiments_subs": [
            (
                "4.1 Experimental Protocol for ML",
                "The synthetic dataset was split 80/20 with stratification on risk_level. Metrics reported "
                "on the test partition. Matching was evaluated qualitatively on demo scenarios with a corpus "
                "of eight seeded attorney profiles aligned to database identifiers. All experiments ran on a "
                "development workstation with Python 3.12, scikit-learn 1.5.2, and FastAPI 0.115.0.",
            ),
            (
                "4.2 Software Integration Tests",
                "Five JUnit suites executed against Testcontainers PostgreSQL: application context load, "
                "ownership enforcement, case sub-resources (timeline and legal responses), feature flows "
                "(password change, status transitions, admin reports), and document upload/download. Frontend "
                "production build (ng build) verified compile-time integrity of physician, attorney, and "
                "admin portals.",
            ),
        ],
        "results_quant_intro": (
            "Table 1 summarizes Random Forest performance on the held-out test set (model version rf-v1)."
        ),
        "table_headers": ["Metric", "Value"],
        "table_data": [
            ("Precision (macro)", "0.6179"),
            ("Recall (macro)", "0.5973"),
            ("F1-score (macro)", "0.6034"),
            ("Dataset size", "6,000 synthetic cases"),
            ("Mean inference latency (local)", "~120 ms"),
        ],
        "table_caption": "Table 1: Risk classifier evaluation metrics.",
        "results_quant_body": (
            "All five backend integration suites passed (100% success) with Docker-enabled PostgreSQL. "
            "Critical REST operations—including login, case creation, matching proxy, and admin "
            "summary—completed under 2 seconds in local measurements, meeting the OE3 response-time "
            "indicator for the MVP scope."
        ),
        "results_qual_title": "5.2 Qualitative Results",
        "results_qual": (
            "End-to-end demonstrations confirmed: (i) role-based navigation across three portals; "
            "(ii) case registration with simulated context; (iii) attorney recommendations with XAI "
            "reasons; (iv) persisted ML classifications retrievable from the API; and (v) visible AI "
            "advisory disclaimers. Attorneys received relevant unassigned cases filtered by declared "
            "medical areas. These observations support feasibility for academic pilot evaluation at "
            "Clínica SANNA \"El Golf\" using simulated scenarios only."
        ),
        "conclusions": [
            "This paper presented Sinapsistencia, a medical-legal mediation platform that integrates "
            "structured workflows, auditability, and explainable machine learning for the Peruvian "
            "health context. The hybrid design—Random Forest for risk scoring and TF-IDF for attorney "
            "matching—addresses cold-start constraints while keeping inference latency suitable for "
            "interactive use.",
            "Results on synthetic data demonstrate moderate classification performance (F1 ≈ 0.60) and "
            "successful software integration across frontend, backend, database, and ML service layers. "
            "The system explicitly treats AI outputs as non-deterministic support, aligning with "
            "ethical requirements for liability-sensitive domains.",
            "Future work includes: (i) expanding the attorney corpus and real pilot feedback without "
            "identifiable patient data; (ii) quantitative matching evaluation toward a ≥70% top-3 "
            "relevance target; (iii) cloud deployment for demo environments; (iv) execution of the full "
            "90-case functional test plan; and (v) usability assessment using the System Usability Scale "
            "in TP2.",
        ],
        "ack": (
            "We thank the School of Systems and Computer Engineering at the Peruvian University of "
            "Applied Sciences for academic support; our advisors Jymmy Stuwart Dextre Alarcon and "
            "Jorge Enrique Polo Martinez for methodological guidance; and Clínica SANNA \"El Golf\" "
            "for institutional collaboration in the simulated validation context."
        ),
        "references": [
            "Aggarwal, C. C. (2016). Recommender systems: The textbook. Springer.",
            "Congreso de la República del Perú. (2011). Ley N.º 29733, Ley de Protección de Datos Personales. El Peruano.",
            "Gunenc, O., Kulhan, N. G., Bayman, M. G., Celik, C., Bilgi, A., Colluoglu, C., & Kulhan, M. (2024). Medical malpractice stress syndrome and defensive medicine in obstetricians and gynecologists in Turkey. International Journal of Clinical Practice, 2024, Article 8226403. https://doi.org/10.1155/2024/8226403",
            "Jones, C., Thornton, J., & Wyatt, J. C. (2023). Artificial intelligence and clinical decision support: Clinicians' perspectives on trust, trustworthiness, and liability. Medical Law Review, 31(4), 501–520. https://doi.org/10.1093/medlaw/fwad013",
            "Manning, C. D., Raghavan, P., & Schütze, H. (2008). Introduction to information retrieval. Cambridge University Press.",
            "Miziara, I. D., & Miziara, C. S. M. G. (2022). Medical errors, medical negligence and defensive medicine: A narrative review. Clinics, 77, Article 100053. https://doi.org/10.1016/j.clinsp.2022.100053",
            "Ricci, F., Rokach, L., & Shapira, B. (2015). Recommender systems handbook (2nd ed.). Springer.",
            "Rowland, S. P., Fitzgerald, J. E., Lungren, M., Lee, E. (H.), Harned, Z., & McGregor, A. H. (2022). Digital health technology-specific risks for medical malpractice liability. NPJ Digital Medicine, 5, Article 157. https://doi.org/10.1038/s41746-022-00698-3",
            "Scott, I. A., van der Vegt, A., Lane, P., McPhail, S., & Magrabi, F. (2024). Achieving large-scale clinician adoption of AI-enabled decision support. BMJ Health & Care Informatics, 31, Article e100971. https://doi.org/10.1136/bmjhci-2023-100971",
            "Superintendencia Nacional de Salud. (2025). SUSALUD brindó más de 163 mil atenciones sobre derechos en salud durante el 2024. Gobierno del Perú.",
            "Tumelty, M.-E. (2023). Plaintiff aims in medical negligence disputes: Limitations of an adversarial system. Medical Law Review, 31(2), 226–246. https://doi.org/10.1093/medlaw/fwac037",
            "Wang, L., Zhang, Z., Wang, D., Cao, W., Zhou, X., Zhang, P., Liu, J., Fan, X., & Tian, F. (2023). Human-centered design and evaluation of AI-empowered clinical decision support systems: A systematic review. Frontiers in Computer Science, 5, Article 1187299. https://doi.org/10.3389/fcomp.2023.1187299",
            "Wong, K. E., Parikh, P. D., Miller, K. C., & Zonfrillo, M. R. (2021). Emergency department and urgent care medical malpractice claims 2001–15. Western Journal of Emergency Medicine, 22(2), 333–338. https://doi.org/10.5811/westjem.2020.9.48845",
        ],
        "pdf_fallback_title": "Sinapsistencia Short Paper TP1",
        "pdf_fallback_body": (
            "Please open Sinapsistencia_ShortPaper_TP1.docx for the full formatted paper. "
            "Automatic PDF conversion requires Microsoft Word installed."
        ),
    },
    "es": {
        "suffix": "_ES",
        "title": (
            "Sinapsistencia: plataforma web de mediación médico-legal con Machine Learning "
            "explicable en el sector salud del Perú"
        ),
        "affiliation": (
            "\u00b9Escuela de Ingenier\u00eda de Sistemas y Computaci\u00f3n, Universidad Peruana de "
            "Ciencias Aplicadas, Av. Primavera 2390, Lima, Per\u00fa"
        ),
        "keywords_label": "Palabras clave:",
        "keywords": (
            "Mediación médico-legal, Machine Learning, Inteligencia artificial explicable, "
            "Random Forest, TF-IDF, Informática en salud."
        ),
        "abstract_label": "Resumen:",
        "abstract": (
            "Los profesionales de la salud en el Perú enfrentan una exposición creciente a reclamos por "
            "mala praxis y carecen con frecuencia de acceso estructurado y oportuno a asesoría legal "
            "especializada. Este trabajo presenta Sinapsistencia, una plataforma web que conecta médicos "
            "con abogados de derecho médico mediante microservicios (Angular 21, Spring Boot 3.5, "
            "PostgreSQL y FastAPI). El sistema registra consultas médico-legales con contexto clínico "
            "simulado conforme a la Ley N.º 29733, ofrece trazabilidad documental, paneles por rol y dos "
            "componentes de ML explicable: Random Forest para riesgo legal y TF-IDF con similitud del "
            "coseno para recomendación de abogados. El modelo de riesgo se entrenó con 6 000 casos "
            "sintéticos y alcanzó precisión macro de 0,618, recall de 0,597 y F1 de 0,603. La validación "
            "incluyó cinco suites automatizadas de integración y escenarios de aceptación incrementales. "
            "La latencia media de inferencia se mantuvo por debajo de 200 ms. Los resultados muestran que "
            "un enfoque híbrido con factores XAI y respaldos de arranque en frío es viable para un MVP "
            "académico en un dominio donde la IA debe actuar como apoyo a la decisión."
        ),
        "sections": {
            "intro": "1 INTRODUCCIÓN",
            "related": "2 TRABAJOS RELACIONADOS",
            "solution": "3 SOLUCIÓN PROPUESTA",
            "experiments": "4 EXPERIMENTOS Y VALIDACIÓN",
            "results": "5 RESULTADOS",
            "conclusions": "6 CONCLUSIONES",
            "ack": "AGRADECIMIENTOS",
            "refs": "REFERENCIAS",
        },
        "intro": [
            "Los conflictos por mala praxis y negligencia médica en el Perú son cada vez más visibles "
            "a través de la supervisión institucional y los canales de defensa de derechos en salud. "
            "SUSALUD reportó más de 163 000 atenciones sobre derechos en salud solo en 2024, lo que "
            "refleja una demanda sostenida de mecanismos formales de recurso. Sin embargo, los médicos "
            "recurren con frecuencia a redes informales para obtener orientación legal, con demoras "
            "estimadas de 48 a 72 horas y escasa trazabilidad auditable (Miziara y Miziara, 2022; "
            "Tumelty, 2023).",
            "El problema es relevante porque la demora en la orientación legal incrementa la vulnerabilidad "
            "profesional, favorece la medicina defensiva y dificulta la estructuración temprana del caso "
            "(Gunenc et al., 2024; Rowland et al., 2022). Las plataformas digitales podrían mejorar la "
            "mediación, pero los enfoques ingenuos fallan por tres razones: (i) los datos médico-legales "
            "son sensibles y regulados; (ii) las recomendaciones de caja negra son éticamente inaceptables "
            "en contextos de responsabilidad (Jones et al., 2023; Scott et al., 2024); y (iii) las "
            "condiciones de arranque en frío limitan el filtrado colaborativo cuando faltan historiales "
            "densos de interacción médico-abogado.",
            "Las soluciones previas abordan fragmentos del problema—historias clínicas electrónicas, "
            "marketplaces legales genéricos o sistemas de apoyo a la decisión clínica—, pero no la "
            "mediación integrada médico-abogado con priorización explicable bajo gobernanza por roles. "
            "Sinapsistencia se diferencia al combinar (1) contexto de caso simulado no identificable, "
            "(2) separación en microservicio de la inferencia ML, (3) persistencia de explicaciones XAI "
            "y (4) emparejamiento de respaldo por especialidad cuando el servicio ML no está disponible.",
            "Los componentes clave del enfoque son: una aplicación web con tres portales (médico, abogado "
            "y administrador); una API Spring Boot con reglas de ownership en la capa de servicio; un "
            "servicio FastAPI que implementa puntuación de riesgo con Random Forest y ranking de abogados "
            "con TF-IDF; y una bitácora de auditoría para acciones críticas. Las limitaciones actuales "
            "incluyen datos de entrenamiento sintéticos, validación piloto sin registros reales de "
            "pacientes y la ejecución pendiente a gran escala del plan académico de 90 casos de prueba.",
            "El resto del artículo se organiza así: la sección 2 revisa trabajos relacionados; la "
            "sección 3 describe la arquitectura y los algoritmos propuestos; la sección 4 explica la "
            "metodología de validación; la sección 5 presenta los resultados; y la sección 6 concluye "
            "con perspectivas futuras.",
        ],
        "rw_intro": (
            "La investigación sobre riesgo médico-legal, apoyo a la decisión asistido por IA y sistemas "
            "de recomendación fundamenta Sinapsistencia. Agrupamos los trabajos previos en tres líneas."
        ),
        "related": [
            (
                "Negligencia médica y responsabilidad digital.",
                "Miziara y Miziara (2022) revisan errores médicos, negligencia y medicina defensiva, "
                "destacando la documentación y la comunicación como puntos recurrentes de falla. Rowland "
                "et al. (2022) analizan riesgos de mala praxis en salud digital, enfatizando trazabilidad "
                "y rendición de cuentas. Wong et al. (2021) estudian reclamos en servicios de urgencia, "
                "mostrando que los problemas procedimentales y documentales dominan los patrones de "
                "litigio. Estos estudios justifican la captura estructurada de casos y el versionado "
                "documental, pero no ofrecen flujos integrados de mediación entre clínicos y abogados.",
            ),
            (
                "IA confiable en contextos clínicos y legales.",
                "Jones et al. (2023) reportan preocupaciones de los clínicos sobre confianza y "
                "responsabilidad al usar apoyo a la decisión con IA. Wang et al. (2023) revisan "
                "sistemáticamente el diseño centrado en el humano para sistemas clínicos con IA, "
                "recomendando transparencia y límites explícitos. Scott et al. (2024) señalan barreras "
                "de adopción cuando faltan explicaciones. Nuestra plataforma aborda esta brecha al "
                "persistir factores de riesgo derivados de la importancia de características y al "
                "mostrar mensajes de advertencia no deterministas en la interfaz.",
            ),
            (
                "Técnicas de emparejamiento y modelado de riesgo.",
                "Manning et al. (2008) y Ricci et al. (2015) establecen TF-IDF y la recomendación basada "
                "en contenido como estrategias efectivas de arranque en frío, relevantes para el "
                "emparejamiento médico-abogado sin historiales densos. Los clasificadores Random Forest "
                "permanecen como línea base robusta para puntuación tabular de riesgo con contribuciones "
                "interpretables (Aggarwal, 2016). A diferencia de marketplaces legales previos, "
                "Sinapsistencia acopla ambas técnicas en un flujo específico de salud con datos "
                "simulados alineados a la normativa peruana de privacidad (Congreso de la República del "
                "Perú, 2011).",
            ),
        ],
        "solution_intro": (
            "Sinapsistencia sigue una arquitectura de microservicios desacoplada. El frontend Angular 21 "
            "consume APIs REST de un backend Spring Boot 3.5 (Java 21) protegido con JWT y control de "
            "acceso basado en roles. PostgreSQL 16 almacena los datos operativos y las migraciones Flyway "
            "versionan el esquema. Un servicio FastAPI dedicado aloja los modelos ML, invocado por el "
            "backend mediante un proxy HTTP con tiempo de espera de 5 segundos y degradación controlada."
        ),
        "solution_subs": [
            (
                "3.1 Arquitectura del sistema",
                "El portal médico permite crear consultas con contexto clínico-legal simulado (edad "
                "referencial, área médica, resumen, factores relevantes) sin datos identificables del "
                "paciente. El portal del abogado lista casos pendientes y solicitudes de contacto "
                "pertinentes. El portal administrador gestiona usuarios, bitácora de auditoría, métricas "
                "ML y reportes resumen. Las reglas de ownership impiden que un médico acceda a consultas "
                "de sus pares; se aplican en la capa de servicio y se validan con pruebas de integración "
                "automatizadas.",
            ),
            (
                "3.2 Modelo de evaluación de riesgo (Random Forest)",
                "El módulo de riesgo predice risk_score (0–1) y risk_level (bajo, moderado, alto, crítico) "
                "a partir de características estructuradas: especialidad médica, complejidad del "
                "procedimiento, prioridad, completitud de documentación, consentimiento informado, "
                "quejas previas y días desde el incidente. Los datos de entrenamiento comprenden 6 000 "
                "filas sintéticas generadas por reglas de negocio más ruido gaussiano para evitar "
                "memorización trivial. Un RandomForestClassifier (200 árboles) y un RandomForestRegressor "
                "comparten el mismo espacio de características preprocesadas (one-hot de especialidad, "
                "ordinal de complejidad/prioridad). La explicabilidad combina importancia global de "
                "características con valores normalizados del caso para producir risk_factors mostrados "
                "como contexto de apoyo.",
            ),
            (
                "3.3 Emparejamiento médico-abogado (TF-IDF + similitud del coseno)",
                "Los perfiles de abogados se vectorizan con TfidfVectorizer sobre especialidades legales, "
                "áreas médicas de interés y texto biográfico. Las consultas del médico, construidas a "
                "partir de especialidad y contexto del caso, se comparan mediante similitud del coseno. "
                "Los resultados incluyen content_score, especialidades coincidentes y razones en lenguaje "
                "natural. Cuando el servicio ML no está disponible, el backend aplica un respaldo de "
                "arranque en frío que ordena abogados por coincidencia de área médica.",
            ),
            (
                "3.4 Contrato de API",
                "FastAPI expone GET /health, GET /api/v1/model/info, POST /api/v1/risk-assessment y "
                "POST /api/v1/recommendations. Spring Boot actúa como proxy de estos endpoints y "
                "normaliza las respuestas a camelCase para el frontend. Los niveles de riesgo alto o "
                "crítico pueden disparar notificaciones opcionales fire-and-forget hacia un webhook "
                "de automatización externo.",
            ),
        ],
        "experiments_intro": (
            "La validación siguió cuatro frentes alineados al plan de calidad del proyecto: (1) evaluación "
            "offline del ML, (2) pruebas de integración del backend, (3) verificación de compilación del "
            "frontend y (4) escenarios funcionales incrementales del catálogo académico de pruebas."
        ),
        "experiments_subs": [
            (
                "4.1 Protocolo experimental para ML",
                "El conjunto sintético se dividió 80/20 con estratificación sobre risk_level. Las métricas "
                "se reportan sobre la partición de prueba. El emparejamiento se evaluó cualitativamente "
                "en escenarios demo con un corpus de ocho perfiles de abogados alineados a identificadores "
                "de base de datos. Todos los experimentos se ejecutaron en una estación de desarrollo con "
                "Python 3.12, scikit-learn 1.5.2 y FastAPI 0.115.0.",
            ),
            (
                "4.2 Pruebas de integración de software",
                "Cinco suites JUnit se ejecutaron contra PostgreSQL con Testcontainers: arranque de "
                "contexto, enforcement de ownership, subrecursos de consulta (timeline y respuestas "
                "legales), flujos de funcionalidades (cambio de contraseña, transiciones de estado, "
                "reportes admin) y carga/descarga de documentos. La compilación de producción del frontend "
                "(ng build) verificó la integridad de los portales médico, abogado y administrador.",
            ),
        ],
        "results_quant_intro": (
            "La Tabla 1 resume el desempeño del Random Forest en el conjunto de prueba retenido "
            "(versión del modelo rf-v1)."
        ),
        "table_headers": ["Métrica", "Valor"],
        "table_data": [
            ("Precisión (macro)", "0,6179"),
            ("Recall (macro)", "0,5973"),
            ("F1-score (macro)", "0,6034"),
            ("Tamaño del dataset", "6 000 casos sintéticos"),
            ("Latencia media de inferencia (local)", "~120 ms"),
        ],
        "table_caption": "Tabla 1: Métricas de evaluación del clasificador de riesgo.",
        "results_quant_body": (
            "Las cinco suites de integración del backend pasaron con éxito del 100 % usando PostgreSQL "
            "con Docker. Las operaciones REST críticas—incluidos login, creación de consulta, proxy de "
            "matching y reporte administrativo—completaron en menos de 2 segundos en mediciones locales, "
            "cumpliendo el indicador de tiempo de respuesta del OE3 para el alcance del MVP."
        ),
        "results_qual_title": "5.2 Resultados cualitativos",
        "results_qual": (
            "Las demostraciones de extremo a extremo confirmaron: (i) navegación por rol en los tres "
            "portales; (ii) registro de consultas con contexto simulado; (iii) recomendaciones de "
            "abogados con razones XAI; (iv) clasificaciones ML persistidas recuperables desde la API; "
            "y (v) advertencias visibles sobre el uso de IA. Los abogados recibieron casos no asignados "
            "pertinentes filtrados por áreas médicas declaradas. Estas observaciones respaldan la "
            "viabilidad de una evaluación piloto académica en la Clínica SANNA «El Golf» usando únicamente "
            "escenarios simulados."
        ),
        "conclusions": [
            "Este artículo presentó Sinapsistencia, una plataforma de mediación médico-legal que integra "
            "flujos estructurados, auditabilidad y Machine Learning explicable para el contexto de salud "
            "peruano. El diseño híbrido—Random Forest para puntuación de riesgo y TF-IDF para "
            "emparejamiento médico-abogado—aborda restricciones de arranque en frío manteniendo una "
            "latencia de inferencia adecuada para uso interactivo.",
            "Los resultados sobre datos sintéticos demuestran un desempeño de clasificación moderado "
            "(F1 ≈ 0,60) e integración exitosa del software entre frontend, backend, base de datos y "
            "servicio ML. El sistema trata explícitamente las salidas de IA como apoyo no determinista, "
            "alineado a requisitos éticos en dominios sensibles a la responsabilidad.",
            "El trabajo futuro incluye: (i) ampliar el corpus de abogados y retroalimentación piloto sin "
            "datos identificables; (ii) evaluación cuantitativa del emparejamiento hacia una meta de "
            "pertinencia ≥ 70 % en el top-3; (iii) despliegue en la nube para entornos de demostración; "
            "(iv) ejecución del plan completo de 90 casos de prueba funcionales; y (v) evaluación de "
            "usabilidad con la Escala de Usabilidad del Sistema (SUS) en TP2.",
        ],
        "ack": (
            "Agradecemos a la Escuela de Ingeniería de Sistemas y Computación de la Universidad "
            "Peruana de Ciencias Aplicadas por el apoyo académico; a nuestros asesores Jymmy Stuwart "
            "Dextre Alarcón y Jorge Enrique Polo Martínez por la orientación metodológica; y a la "
            "Clínica SANNA «El Golf» por la colaboración institucional en el contexto de validación "
            "simulada."
        ),
        "references": [
            "Aggarwal, C. C. (2016). Recommender systems: The textbook. Springer.",
            "Congreso de la República del Perú. (2011). Ley N.º 29733, Ley de Protección de Datos Personales. El Peruano.",
            "Gunenc, O., Kulhan, N. G., Bayman, M. G., Celik, C., Bilgi, A., Colluoglu, C., & Kulhan, M. (2024). Medical malpractice stress syndrome and defensive medicine in obstetricians and gynecologists in Turkey. International Journal of Clinical Practice, 2024, Article 8226403. https://doi.org/10.1155/2024/8226403",
            "Jones, C., Thornton, J., & Wyatt, J. C. (2023). Artificial intelligence and clinical decision support: Clinicians' perspectives on trust, trustworthiness, and liability. Medical Law Review, 31(4), 501–520. https://doi.org/10.1093/medlaw/fwad013",
            "Manning, C. D., Raghavan, P., & Schütze, H. (2008). Introduction to information retrieval. Cambridge University Press.",
            "Miziara, I. D., & Miziara, C. S. M. G. (2022). Medical errors, medical negligence and defensive medicine: A narrative review. Clinics, 77, Article 100053. https://doi.org/10.1016/j.clinsp.2022.100053",
            "Ricci, F., Rokach, L., & Shapira, B. (2015). Recommender systems handbook (2nd ed.). Springer.",
            "Rowland, S. P., Fitzgerald, J. E., Lungren, M., Lee, E. (H.), Harned, Z., & McGregor, A. H. (2022). Digital health technology-specific risks for medical malpractice liability. NPJ Digital Medicine, 5, Article 157. https://doi.org/10.1038/s41746-022-00698-3",
            "Scott, I. A., van der Vegt, A., Lane, P., McPhail, S., & Magrabi, F. (2024). Achieving large-scale clinician adoption of AI-enabled decision support. BMJ Health & Care Informatics, 31, Article e100971. https://doi.org/10.1136/bmjhci-2023-100971",
            "Superintendencia Nacional de Salud. (2025). SUSALUD brindó más de 163 mil atenciones sobre derechos en salud durante el 2024. Gobierno del Perú.",
            "Tumelty, M.-E. (2023). Plaintiff aims in medical negligence disputes: Limitations of an adversarial system. Medical Law Review, 31(2), 226–246. https://doi.org/10.1093/medlaw/fwac037",
            "Wang, L., Zhang, Z., Wang, D., Cao, W., Zhou, X., Zhang, P., Liu, J., Fan, X., & Tian, F. (2023). Human-centered design and evaluation of AI-empowered clinical decision support systems: A systematic review. Frontiers in Computer Science, 5, Article 1187299. https://doi.org/10.3389/fcomp.2023.1187299",
            "Wong, K. E., Parikh, P. D., Miller, K. C., & Zonfrillo, M. R. (2021). Emergency department and urgent care medical malpractice claims 2001–15. Western Journal of Emergency Medicine, 22(2), 333–338. https://doi.org/10.5811/westjem.2020.9.48845",
        ],
        "pdf_fallback_title": "Sinapsistencia Short Paper TP1 (ES)",
        "pdf_fallback_body": (
            "Abra Sinapsistencia_ShortPaper_TP1_ES.docx para la versión completa formateada. "
            "La conversión automática a PDF requiere Microsoft Word instalado."
        ),
    },
}


def set_run_font(run, bold=False, italic=False, size=None):
    run.font.name = FONT_NAME
    run.font.size = size or FONT_SIZE
    run.bold = bold
    run.italic = italic


def add_paragraph(doc, text, bold=False, align=None, space_after=6, italic=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, bold=bold, italic=italic)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, bold=True, size=HEADING_SIZE)
    return p


def build_document(lang: str) -> Document:
    c = CONTENT[lang]
    sec = c["sections"]
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(c["title"])
    set_run_font(tr, bold=True, size=Pt(12))

    add_paragraph(
        doc,
        "Augusto Mathias Leonardo Vasquez Requejo\u00b9\u1d43, Renato German Reyes Valenzuela\u00b9\u1d47",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )
    add_paragraph(doc, c["affiliation"], align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph(
        doc,
        "u20221a955@upc.edu.pe, u20221b471@upc.edu.pe",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )

    add_paragraph(doc, f"{c['keywords_label']} {c['keywords']}", bold=True, space_after=8)
    add_paragraph(doc, c["abstract_label"], bold=True, space_after=4)
    add_paragraph(doc, c["abstract"], space_after=10)

    add_heading(doc, sec["intro"])
    for para in c["intro"]:
        add_paragraph(doc, para)

    add_heading(doc, sec["related"])
    add_paragraph(doc, c["rw_intro"])
    for rtitle, body in c["related"]:
        add_paragraph(doc, rtitle, bold=True, space_after=2)
        add_paragraph(doc, body)

    add_heading(doc, sec["solution"])
    add_paragraph(doc, c["solution_intro"])
    for stitle, body in c["solution_subs"]:
        add_paragraph(doc, stitle, bold=True, space_after=2)
        add_paragraph(doc, body)

    add_heading(doc, sec["experiments"])
    add_paragraph(doc, c["experiments_intro"])
    for stitle, body in c["experiments_subs"]:
        add_paragraph(doc, stitle, bold=True, space_after=2)
        add_paragraph(doc, body)

    add_heading(doc, sec["results"])
    qual_title = "5.1 Quantitative Results" if lang == "en" else "5.1 Resultados cuantitativos"
    add_paragraph(doc, qual_title, bold=True, space_after=2)
    add_paragraph(doc, c["results_quant_intro"], space_after=4)

    table = doc.add_table(rows=len(c["table_data"]) + 1, cols=2)
    table.style = "Table Grid"
    for j, h in enumerate(c["table_headers"]):
        set_run_font(table.rows[0].cells[j].paragraphs[0].add_run(h), bold=True)
    for i, (k, v) in enumerate(c["table_data"], start=1):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

    add_paragraph(doc, c["table_caption"], italic=True, space_after=8)
    add_paragraph(doc, c["results_quant_body"], space_after=6)
    add_paragraph(doc, c["results_qual_title"], bold=True, space_after=2)
    add_paragraph(doc, c["results_qual"])

    add_heading(doc, sec["conclusions"])
    for para in c["conclusions"]:
        add_paragraph(doc, para)

    add_heading(doc, sec["ack"])
    add_paragraph(doc, c["ack"])

    add_heading(doc, sec["refs"])
    for ref in c["references"]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(3)
        set_run_font(p.add_run(ref))

    return doc


def export_pdf(docx_path: Path, pdf_path: Path) -> bool:
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        return pdf_path.exists()
    except Exception as exc:
        print(f"docx2pdf failed: {exc}")
        return False


def generate(lang: str) -> None:
    c = CONTENT[lang]
    docx_path = OUT_DIR / f"Sinapsistencia_ShortPaper_TP1{c['suffix']}.docx"
    pdf_path = OUT_DIR / f"Sinapsistencia_ShortPaper_TP1{c['suffix']}.pdf"

    doc = build_document(lang)
    doc.save(docx_path)
    print(f"DOCX ({lang}): {docx_path}")

    if export_pdf(docx_path, pdf_path):
        print(f"PDF  ({lang}): {pdf_path}")
    else:
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

            styles = getSampleStyleSheet()
            pdf = SimpleDocTemplate(str(pdf_path), pagesize=letter)
            story = [
                Paragraph(c["pdf_fallback_title"], styles["Title"]),
                Spacer(1, 12),
                Paragraph(c["pdf_fallback_body"], styles["Normal"]),
            ]
            pdf.build(story)
            print(f"PDF (fallback, {lang}): {pdf_path}")
        except Exception as exc2:
            print(f"PDF generation failed ({lang}): {exc2}")


def main():
    parser = argparse.ArgumentParser(description="Genera Short Paper TP1 Sinapsistencia")
    parser.add_argument(
        "--lang",
        choices=["en", "es"],
        action="append",
        help="Idioma(s) a generar (por defecto: en)",
    )
    args = parser.parse_args()
    langs = args.lang or ["en"]
    for lang in langs:
        generate(lang)


if __name__ == "__main__":
    main()
